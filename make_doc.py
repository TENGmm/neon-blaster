from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 標題
title = doc.add_heading('Neon Blaster - 迷途小怪獸', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('一個使用原生 HTML5 Canvas + ES6 Class 開發的 2D 射擊遊戲\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run2 = info.add_run('適用於學習作品集展示（物件導向程式設計）')
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()
doc.add_heading('📌 遊戲概覽', level=1)
doc.add_paragraph(
    'Neon Blaster 是一款復古霓虹風格的 2D 垂直射擊遊戲，'
    '玩家操控飛船在無盡波次中消滅敵人。遊戲透過升級、COMBO 連擊、'
    '成就系統提供豐富的遊玩體驗，同時展示完整的 ES6 Class 物件導向設計。'
)

doc.add_heading('🎮 遊戲操作', level=1)
ops = [
    ('↑ ↓ ← → 或 WASD', '移動飛船'),
    ('空白鍵', '發射子彈'),
    ('P', '暫停 / 繼續'),
    ('ESC 或 Enter', '返回主選單'),
    ('Tab', '查看成就面板'),
]
table = doc.add_table(rows=len(ops)+1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '按鍵'
hdr[1].text = '功能'
for i, (k, v) in enumerate(ops):
    table.rows[i+1].cells[0].text = k
    table.rows[i+1].cells[1].text = v

doc.add_heading('👾 敵人類型', level=1)
enemies = [
    ('BasicEnemy', '◉', '追蹤玩家，直線前進', '30', '10'),
    ('FastEnemy', '◈', '高速移動，體積小', '20', '20'),
    ('TankEnemy', '◉', '高血量，緩慢，有白色邊框', '80', '30'),
    ('ShooterEnemy', '◎', '會向玩家發射子彈', '40', '25'),
    ('ExploderEnemy', '◌', '接近時會爆炸', '15', '15'),
    ('BossEnemy', '✪', '巨大，高血量，傷害高', '500', '200'),
]
table2 = doc.add_table(rows=len(enemies)+1, cols=5)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = '敵人'
hdr2[1].text = '圖示'
hdr2[2].text = '特性'
hdr2[3].text = '血量'
hdr2[4].text = '分數'
for i, row in enumerate(enemies):
    for j, val in enumerate(row):
        table2.rows[i+1].cells[j].text = val

doc.add_heading('🎁 道具系統', level=1)
powerups = [
    ('🟢 回復', '恢復 30 血量'),
    ('🟠 炸彈', '清除畫面上所有敵人'),
    ('🔵 護盾', '無敵 5 秒'),
    ('🟡 分數', '2 倍分數持續 10 秒'),
    ('🟣 速度', '移動速度 +60%'),
    ('🔷 冰凍', '所有敵人凍結 5 秒'),
    ('🟤 磁鐵', '自動吸引附近道具'),
]
table3 = doc.add_table(rows=len(powerups)+1, cols=2)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = '道具'
hdr3[1].text = '效果'
for i, (k, v) in enumerate(powerups):
    table3.rows[i+1].cells[0].text = k
    table3.rows[i+1].cells[1].text = v

doc.add_heading('🔥 主要系統', level=1)
doc.add_heading('COMBO 連擊系統', level=2)
doc.add_paragraph('2 秒內連續殺敵可累積連擊數，連擊數越高顯示效果越華麗。10 連擊 = 橙色，20 連擊 = 紫色。')

doc.add_heading('成就系統（12 個成就）', level=2)
achievements = [
    ('初試啼聲', '🎯', '殺死第一個敵人'),
    ('小試身手', '🔟', '累計擊殺 10'),
    ('殺敵好手', '💀', '累計擊殺 50'),
    ('百連斬', '⚔️', '累計擊殺 100'),
    ('連擊新手', '🔥', '達成 10 連擊'),
    ('連擊達人', '✨', '達成 25 連擊'),
    ('初窺門徑', '🌟', '升到 5 級'),
    ('略有小成', '⭐', '升到 10 級'),
    ('波瀾不驚', '🌊', '到達第 10 波'),
    ('滴水不漏', '🛡️', '無傷通關'),
    ('Boss殺手', '👹', '擊敗 Boss'),
    ('高分選手', '🏆', '分數超過 5000'),
]
ach = doc.add_table(rows=len(achievements)+1, cols=3)
ach.style = 'Table Grid'
ach.rows[0].cells[0].text = '成就'
ach.rows[0].cells[1].text = '圖示'
ach.rows[0].cells[2].text = '條件'
for i, (name, icon, cond) in enumerate(achievements):
    ach.rows[i+1].cells[0].text = name
    ach.rows[i+1].cells[1].text = icon
    ach.rows[i+1].cells[2].text = cond

doc.add_heading('經驗值 / 升級系統', level=2)
doc.add_paragraph('殺敵獲得經驗值，升級提升屬性：回血、加速、射速加快、分數倍率提升。')

doc.add_heading('結算報告', level=2)
doc.add_paragraph('遊戲結束顯示完整戰鬥報告：評級（S/A/B/C/D）、分數、擊殺數、最高連擊、存活時間、波次、等級、命中率。')

doc.add_heading('🏗️ 技術架構', level=1)
doc.add_paragraph('本遊戲以 ES6 Class 實作物件導向設計，展示以下 OOP 概念：')

oop = [
    ('封裝（Encapsulation）', '屬性與方法包裝在類別內'),
    ('繼承（Inheritance）', '子類別繼承 GameObject 基礎類'),
    ('多形（Polymorphism）', '不同敵人類別覆寫 draw() 和 update() 方法'),
    ('工廠模式（Factory Pattern）', 'EnemyFactory.create() 統一建立敵人'),
]
for concept, desc in oop:
    p = doc.add_paragraph()
    run = p.add_run(f'• {concept}：')
    run.bold = True
    p.add_run(desc)

doc.add_heading('類別層級', level=2)
classes = """GameObject（基礎類）
├── Player（玩家）
├── Bullet（子彈）
├── Enemy（敵人）
│   ├── BasicEnemy / FastEnemy / TankEnemy
│   ├── ShooterEnemy / ExploderEnemy / BossEnemy
├── PowerUp（道具）
│   ├── HealthUp / BombUp / ShieldUp / ScoreUp / SpeedUp / FreezeUp / MagnetUp
├── Particle（粒子效果）
└── FloatingText（浮動文字）"""
p = doc.add_paragraph()
run = p.add_run(classes)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('🎨 視覺效果', level=1)
effects = [
    '霓虹風格（Neon Glow Effect）',
    '粒子爆炸效果',
    '子彈軌跡',
    '敵人受擊閃爍',
    '畫面震動效果',
    '背景網格動畫',
]
for e in effects:
    doc.add_paragraph(f'• {e}', style='List Bullet')

doc.add_heading('🚀 部署方式', level=1)
doc.add_paragraph('直接用瀏覽器開啟 index.html 即可遊玩，無需任何伺服器或安裝。')
doc.add_paragraph('GitHub：https://github.com/TENGmm/neon-blaster')

doc.save('/Users/teng/Desktop/WebGame/Neon_Blaster_介紹.docx')
print('done')
